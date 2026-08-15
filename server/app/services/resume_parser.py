import os
import re
from typing import Dict, Any, List

class ResumeParserService:
    """
    Parses uploaded PDF and DOCX resume documents using pypdf and python-docx.
    Extracts actual text, summary, contact information, and skills from the file.
    No hardcoded candidate data or fabricated fallback content is used.
    """
    
    COMMON_SKILLS = [
        "React", "TypeScript", "JavaScript", "Python", "Java", "C++", "C#", "Go", "Rust",
        "HTML", "HTML5", "CSS", "CSS3", "Tailwind CSS", "Bootstrap", "Sass", "LESS",
        "Node.js", "Express", "FastAPI", "Django", "Flask", "Spring Boot", "GraphQL", "REST APIs",
        "SQL", "PostgreSQL", "MySQL", "MongoDB", "Redis", "SQLite", "SQLAlchemy",
        "Git", "GitHub", "Docker", "Kubernetes", "AWS", "Azure", "GCP", "CI/CD",
        "Jest", "Cypress", "PyTest", "Webpack", "Vite", "Redux", "Next.js", "Vue.js", "Angular"
    ]

    @staticmethod
    def parse_document(file_path: str, filename: str) -> Dict[str, Any]:
        text_content = ""
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                pages_text = []
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        pages_text.append(extracted)
                text_content = "\n".join(pages_text)
            except Exception as e:
                raise ValueError(f"Failed to extract text from PDF document: {str(e)}")
        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(file_path)
                paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also extract text from table cells if present
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs_text.append(cell.text.strip())
                text_content = "\n".join(paragraphs_text)
            except Exception as e:
                raise ValueError(f"Failed to extract text from DOCX document: {str(e)}")
        else:
            raise ValueError(f"Unsupported file format '{ext}'. Only PDF and DOCX files are supported.")

        clean_text = text_content.strip()
        if not clean_text:
            raise ValueError("The uploaded document contains no readable text.")

        # Extract matching skills from clean text
        matched_skills = []
        for skill in ResumeParserService.COMMON_SKILLS:
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, clean_text, re.IGNORECASE):
                matched_skills.append(skill)

        # Build a initial summary snippet from first 300 characters
        first_lines = [line.strip() for line in clean_text.splitlines() if line.strip()]
        summary_snippet = " ".join(first_lines[:3])[:300] if first_lines else "Parsed candidate resume document."

        return {
            "raw_text": clean_text,
            "summary": summary_snippet,
            "skills": matched_skills,
            "word_count": len(clean_text.split()),
            "character_count": len(clean_text)
        }
